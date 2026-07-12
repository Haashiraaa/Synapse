# src/telegram/media.py


import base64
import io
from dataclasses import dataclass
from typing import Any

import docx

from src.exceptions.errors import UnsupportedMediaError


@dataclass
class MediaOutcome:
    placeholder: str  # what gets persisted to messages.content
    # base64 vision block, or None for text-embedded media
    media_block: dict[str, Any] | None = None


class MediaProcessor:
    """
    Format-specific media handling. Add a new type by adding one method
    here + one branch in process()/supports() — handlers.py never changes.
    """

    IMAGE_TYPES = {"image/jpeg", "image/png", "image/webp", "image/gif"}
    DIFF_EXTENSIONS = (".diff", ".patch")
    DOCX_EXTENSIONS = (".docx",)
    MAX_DIFF_BYTES = 250 * 1024
    MAX_DOCX_CHARS = 20_000

    def supports(self, mime_type: str | None, filename: str | None) -> bool:
        if mime_type in self.IMAGE_TYPES or mime_type == "application/pdf":
            return True
        if filename and filename.lower().endswith(self.DIFF_EXTENSIONS + self.DOCX_EXTENSIONS):
            return True
        return False

    def process(
        self, raw: bytes, mime_type: str | None, filename: str | None, caption: str
    ) -> MediaOutcome:
        if filename and filename.lower().endswith(self.DIFF_EXTENSIONS):
            return self._handle_diff(raw, filename, caption)
        if filename and filename.lower().endswith(self.DOCX_EXTENSIONS):
            return self._handle_docx(raw, filename, caption)
        if mime_type == "application/pdf":
            return self._handle_pdf(raw, caption)
        if mime_type in self.IMAGE_TYPES:
            return self._handle_image(raw, mime_type, caption)
        raise UnsupportedMediaError(
            f"Unsupported media: mime={mime_type} filename={filename}")

    # ── per-type handlers ─────────────────────────────────────

    def _handle_image(self, raw: bytes, mime_type: str, caption: str) -> MediaOutcome:
        b64 = base64.b64encode(raw).decode("utf-8")
        return MediaOutcome(
            placeholder="[sent an image]" +
            (f": {caption}" if caption else ""),
            media_block={
                "type": "image",
                "source": {"type": "base64", "media_type": mime_type, "data": b64},
            },
        )

    def _handle_pdf(self, raw: bytes, caption: str) -> MediaOutcome:
        b64 = base64.b64encode(raw).decode("utf-8")
        return MediaOutcome(
            placeholder="[sent a PDF]" + (f": {caption}" if caption else ""),
            media_block={
                "type": "document",
                "source": {"type": "base64", "media_type": "application/pdf", "data": b64},
            },
        )

    def _handle_diff(self, raw: bytes, filename: str, caption: str) -> MediaOutcome:
        text = raw.decode("utf-8", errors="replace")
        truncated = len(text.encode("utf-8")) > self.MAX_DIFF_BYTES
        if truncated:
            text = text[: self.MAX_DIFF_BYTES]
        body = f"[sent diff: {filename}]" + \
            (f" — {caption}" if caption else "")
        body += f"\n```diff\n{text}\n```"
        if truncated:
            body += "\n[diff truncated — too large]"
        return MediaOutcome(placeholder=body)  # no media_block → text pipeline

    def _handle_docx(self, raw: bytes, filename: str, caption: str) -> MediaOutcome:
        try:
            document = docx.Document(io.BytesIO(raw))
        except Exception:
            return MediaOutcome(
                placeholder=f"[sent {filename}, but I couldn't read it — "
                "the file may be corrupted or not a valid .docx]"
            )

        parts = [p.text for p in document.paragraphs if p.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells]
                if any(cells):
                    parts.append(" | ".join(cells))
        text = "\n".join(parts)

        truncated = len(text) > self.MAX_DOCX_CHARS
        if truncated:
            text = text[: self.MAX_DOCX_CHARS]
        if not text.strip():
            text = "(no extractable text — file may be images, headers/footers only, or empty)"

        body = f"[sent docx: {filename}]" + \
            (f" — {caption}" if caption else "")
        body += f"\n{text}"
        if truncated:
            body += "\n[docx truncated — too large]"
        return MediaOutcome(placeholder=body)  # no media_block → text pipeline
